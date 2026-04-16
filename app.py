import os
# These MUST be set before importing torch/faiss/sentence_transformers
# to prevent segfaults from OpenMP conflicts on macOS
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["OMP_NUM_THREADS"] = "1"

import streamlit as st
import faiss
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings, HuggingFaceEndpoint, ChatHuggingFace
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore


def process_input(input_type, input_data):
    """Processes different input types and returns a vectorstore."""
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif hasattr(input_data, 'read'):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for PDF")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data  # Input is already a text string
        else:
            raise ValueError("Expected a string for 'Text' input type.")
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif hasattr(input_data, 'read'):
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for DOCX")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif hasattr(input_data, 'read'):
            text = str(input_data.read().decode('utf-8'))
        else:
            raise ValueError("Invalid input data for TXT")
        documents = text
    else:
        raise ValueError("Unsupported input type")

    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]  # Access page_content from each Document
    else:
        texts = text_splitter.split_text(documents)

    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}

    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )
    # Create FAISS index
    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)
    # Create FAISS vector store with the embedding function
    vector_store = FAISS(
        embedding_function=hf_embeddings,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)  # Add documents to the vector store
    return vector_store


def answer_question(vectorstore, query, api_key):
    """Answers a question based on the provided vectorstore."""
    # Use ChatHuggingFace since Llama-3 is a conversational model,
    # not a plain text-generation model on the HF Inference API.
    llm = HuggingFaceEndpoint(
        repo_id='meta-llama/Meta-Llama-3-8B-Instruct',
        huggingfacehub_api_token=api_key,
        temperature=0.6,
    )
    chat_model = ChatHuggingFace(llm=llm)

    qa = RetrievalQA.from_chain_type(llm=chat_model, retriever=vectorstore.as_retriever())
    answer = qa.invoke({"query": query})
    return answer


def main():
    st.title("RAG Q&A App")

    # --- API Key input in the sidebar ---
    api_key = st.sidebar.text_input(
        "🔑 Hugging Face API Key",
        type="password",
        placeholder="hf_...",
        help="Get your free API key from https://huggingface.co/settings/tokens. Make sure to turn on 'Make calls to the Serverless Inference API' permissions!"
    )

    if api_key:
        os.environ['HUGGINGFACEHUB_API_TOKEN'] = api_key
        st.sidebar.success("API key set ✅")
    else:
        st.sidebar.warning("Please enter your Hugging Face API key to proceed.")

    # --- Input type selection ---
    input_type = st.selectbox("Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])

    if input_type == "Link":
        number_input = st.number_input(
            min_value=1, max_value=20, step=1,
            label="Enter the number of Links"
        )
        input_data = []
        for i in range(number_input):
            url = st.sidebar.text_input(f"URL {i+1}")
            input_data.append(url)
    elif input_type == "Text":
        input_data = st.text_input("Enter the text")
    elif input_type == 'PDF':
        input_data = st.file_uploader("Upload a PDF file", type=["pdf"])
    elif input_type == 'TXT':
        input_data = st.file_uploader("Upload a text file", type=['txt'])
    elif input_type == 'DOCX':
        input_data = st.file_uploader("Upload a DOCX file", type=['docx', 'doc'])

    # --- Process input ---
    if st.button("Proceed"):
        if not api_key:
            st.error("⚠️ Please enter your Hugging Face API key in the sidebar first.")
        else:
            with st.spinner("Processing your input..."):
                vectorstore = process_input(input_type, input_data)
                st.session_state["vectorstore"] = vectorstore
            st.success("✅ Input processed! You can now ask questions below.")

    # --- Q&A ---
    if "vectorstore" in st.session_state:
        query = st.text_input("Ask your question")
        if st.button("Submit"):
            if not api_key:
                st.error("⚠️ Please enter your Hugging Face API key in the sidebar first.")
            elif not query:
                st.warning("Please enter a question.")
            else:
                with st.spinner("Generating answer..."):
                    answer = answer_question(st.session_state["vectorstore"], query, api_key)
                    st.write(answer["result"])


if __name__ == "__main__":
    main()